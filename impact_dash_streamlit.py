import streamlit as st
import pandas as pd
import numpy as np
import pickle
import json
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import statsmodels.api as sm
from scipy import stats
from scipy.stats import beta, norm
from datetime import datetime
import math
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
import io

# Page configuration
st.set_page_config(
    page_title="ImPACT Normative Classification",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional styling
st.markdown("""
<style>
.main-header {
    background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
    color: white;
    padding: 20px;
    border-radius: 10px;
    margin-bottom: 20px;
    text-align: center;
}

.metric-card {
    background: white;
    padding: 20px;
    border-radius: 10px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    border-left: 4px solid #2a5298;
    margin: 10px 0;
    text-align: center;
}

.classification-card {
    padding: 20px;
    border-radius: 15px;
    margin: 15px 0;
    text-align: center;
    box-shadow: 0 4px 8px rgba(0,0,0,0.1);
}

.compute-button {
    background: #2a5298;
    color: white;
    border: none;
    padding: 15px 30px;
    border-radius: 8px;
    font-size: 16px;
    font-weight: bold;
    cursor: pointer;
    width: 100%;
    margin: 20px 0;
}

.center-text {
    text-align: center;
}

.stButton > button {
    background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
    color: white;
    border: none;
    border-radius: 8px;
    padding: 12px 24px;
    font-weight: bold;
    width: 100%;
}
</style>
""", unsafe_allow_html=True)

# Load models and parameters
@st.cache_data
def load_models_and_parameters():
    with open('cognitive_models.pkl', 'rb') as f:
        models = pickle.load(f)
    
    with open('model_parameters.json', 'r') as f:
        params = json.load(f)
    
    return models, params

def adjust_single_score(raw_score, age, sex, education, model, adj_mean, is_reaction_speed=False):
    """Adjust a single score for demographics using OLS model (for ES method)"""
    sex_encoded = 1 if sex == 'm' else 0
    
    if is_reaction_speed:
        score_to_adjust = 1 / raw_score
    else:
        score_to_adjust = raw_score
    
    X = pd.DataFrame({
        'const': [1],
        'age': [age],
        'sex_encoded': [sex_encoded],
        'education': [education]
    })
    
    pred = model.predict(X)[0]
    adj_score = score_to_adjust - (pred - adj_mean)
    
    return adj_score

def assign_es_category(adj_score, thresholds):
    """Assign ES category based on adjusted score and thresholds using R methodology"""
    # Check if this uses the R methodology (has observation numbers) or fallback percentile method
    if 'oTL_obs' in thresholds:
        # R methodology - use the proper threshold values
        if adj_score <= thresholds['oTL']:
            return 0
        elif adj_score <= thresholds['threshold_1']:
            return 1
        elif adj_score <= thresholds['threshold_2']:
            return 2
        elif adj_score < thresholds['threshold_3']:
            return 3
        else:
            return 4
    else:
        # Fallback percentile method
        if adj_score <= thresholds['oTL']:
            return 0
        elif adj_score <= thresholds['threshold_1']:
            return 1
        elif adj_score <= thresholds['threshold_2']:
            return 2
        elif adj_score < thresholds['threshold_3']:
            return 3
        else:
            return 4

def get_es_interpretation(es_score):
    """Provide clinical interpretation of ES scores"""
    interpretations = {
        0: {"label": "Impaired ", "color": "#d32f2f", "description":''},
        1: {"label": "Borderline", "color": "#f57c00", "description": ''},
        2: {"label": "Low Normal", "color": "#fbc02d", "description": ''},
        3: {"label": "Normal", "color": "#689f38", "description": ''},
        4: {"label": "High Normal", "color": "#388e3c", "description": ''}
    }
    return interpretations[es_score]

def compute_iverson_percentile(raw_score, age, sex, education, model, mse_resid, is_reaction_speed=False):
    """Compute percentile using Iverson methodology"""
    sex_encoded = 1 if sex == 'm' else 0
    
    if is_reaction_speed:
        score_to_analyze = 1 / raw_score
    else:
        score_to_analyze = raw_score
    
    X = pd.DataFrame({
        'const': [1],
        'age': [age],
        'sex_encoded': [sex_encoded],
        'education': [education]
    })
    
    y_pred = model.predict(X)[0]
    residual = score_to_analyze - y_pred
    sigma = np.sqrt(mse_resid)
    z_score = residual / sigma
    percentile = stats.norm.cdf(z_score) * 100
    percentile = max(1.0, min(99.0, percentile))
    
    return percentile

def compute_iverson_classification(percentiles):
    """Apply Iverson collective classification rules"""
    below_25th = sum(p <= 25 for p in percentiles)
    below_16th = sum(p <= 16 for p in percentiles)
    below_10th = sum(p <= 10 for p in percentiles)
    below_5th = sum(p <= 5 for p in percentiles)
    below_2nd = sum(p <= 2 for p in percentiles)
    
    if below_5th >= 3 or below_2nd >= 2:
        return "Extremely Low"
    elif below_10th >= 3 or below_5th >= 2 or any(p <= 2 for p in percentiles):
        return "Unusually Low"
    elif below_16th >= 3 or below_10th >= 2 or any(p <= 5 for p in percentiles):
        return "Well Below Average"
    elif below_25th >= 3 or below_16th >= 2 or any(p <= 10 for p in percentiles):
        return "Below Average"
    else:
        return "Broadly Normal"

def get_iverson_interpretation(classification):
    """Provide clinical interpretation for Iverson classifications"""
    interpretations = {
        "Broadly Normal": {"color": "#388e3c", "description": "Cognitive performance within normal limits", "clinical_note": "No significant cognitive concerns identified"},
        "Below Average": {"color": "#fbc02d", "description": "Mild cognitive difficulties across domains", "clinical_note": "Consider monitoring and supportive interventions"},
        "Well Below Average": {"color": "#f57c00", "description": "Moderate cognitive difficulties requiring attention", "clinical_note": "Comprehensive assessment and intervention recommended"},
        "Unusually Low": {"color": "#e65100", "description": "Significant cognitive impairment across domains", "clinical_note": "Immediate clinical attention and specialized assessment needed"},
        "Extremely Low": {"color": "#d32f2f", "description": "Severe cognitive impairment requiring urgent intervention", "clinical_note": "Urgent referral for specialized neuropsychological evaluation"}
    }
    return interpretations[classification]

def create_gauge_chart(value, title, max_value=100, color_ranges=None):
    """Create a professional gauge chart for percentiles"""
    if color_ranges is None:
        color_ranges = [
            (0, 5, "#d32f2f"),      # Very low
            (5, 10, "#e65100"),     # Low
            (10, 16, "#f57c00"),    # Below average
            (16, 25, "#fbc02d"),    # Low average
            (25, 100, "#388e3c")    # Normal+
        ]
    
    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = value,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': title, 'font': {'size': 16, 'color': '#1e3c72'}},
        number = {'suffix': "%", 'font': {'size': 20, 'color': '#1e3c72'}},
        gauge = {
            'axis': {'range': [None, max_value], 'tickwidth': 1, 'tickcolor': "darkblue"},
            'bar': {'color': "darkblue"},
            'bgcolor': "white",
            'borderwidth': 2,
            'bordercolor': "gray",
            'steps': [{'range': [start, end], 'color': color} for start, end, color in color_ranges],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': value
            }
        }
    ))
    
    fig.update_layout(height=300, margin=dict(l=20, r=20, t=40, b=20))
    return fig

def create_es_gauge_chart(value, title):
    """Create a professional gauge chart for ES scores (0-4)"""
    color_ranges = [
        (0, 1, "#d32f2f"),      # ES 0
        (1, 2, "#f57c00"),      # ES 1
        (2, 3, "#fbc02d"),      # ES 2
        (3, 4, "#689f38"),      # ES 3
        (4, 5, "#388e3c")       # ES 4
    ]
    
    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = value,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': title, 'font': {'size': 16, 'color': '#1e3c72'}},
        number = {'font': {'size': 24, 'color': '#1e3c72'}},
        gauge = {
            'axis': {'range': [0, 4], 'tickwidth': 1, 'tickcolor': "darkblue", 'tickmode': 'linear', 'tick0': 0, 'dtick': 1},
            'bar': {'color': "darkblue"},
            'bgcolor': "white",
            'borderwidth': 2,
            'bordercolor': "gray",
            'steps': [{'range': [start, end], 'color': color} for start, end, color in color_ranges],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': value
            }
        }
    ))
    
    fig.update_layout(height=300, margin=dict(l=20, r=20, t=45, b=20))
    return fig

def generate_word_report(method, patient_data, results_data):
    """Generate comprehensive assessment report in Word format with both ES and Iverson tables"""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(16)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Header style
    header_style = styles.add_style('CustomHeader', WD_STYLE_TYPE.PARAGRAPH)
    header_font = header_style.font
    header_font.name = 'Arial'
    header_font.size = Pt(14)
    header_font.bold = True
    header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_style.paragraph_format.space_before = Pt(12)
    header_style.paragraph_format.space_after = Pt(6)
    
    # Body style
    body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
    body_font = body_style.font
    body_font.name = 'Arial'
    body_font.size = Pt(11)
    body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body_style.paragraph_format.space_after = Pt(6)
    
    # Add content
    title = doc.add_paragraph('ImPACT NORMATIVE CLASSIFICATION REPORT', style='CustomTitle')
    
    # Subject Information
    doc.add_paragraph('Subject Information', style='CustomHeader')
    doc.add_paragraph(f'Age - {patient_data["age"]} years', style='CustomBody')
    doc.add_paragraph(f'Sex - {"Male" if patient_data["sex"] == "m" else "Female"}', style='CustomBody')
    doc.add_paragraph(f'Education - {patient_data["education"]} years', style='CustomBody')
    doc.add_paragraph(f'Assessment Date - {datetime.now().strftime("%B %d, %Y")}', style='CustomBody')
    
    # Test Results Summary Table
    doc.add_paragraph('Test Results Summary', style='CustomHeader')
    
    # Create comprehensive table (always 5 columns to include both methods)
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    
    # Header row with enhanced formatting
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Test'
    header_cells[1].text = 'Raw Score'
    header_cells[2].text = 'Adjusted Score (ES)' if method == "Individual Outcome Classification (ES Method)" else 'Estimated Percentile'
    header_cells[3].text = 'Individual ES Classification'
    header_cells[4].text = 'Iverson & Schaltz Criteria'
    
    # Format headers with bold and center alignment
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
        # Add shading to header using a more compatible method
        try:
            from docx.oxml import parse_xml
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            cell._element.get_or_add_tcPr().append(shading_elm)
        except:
            # If shading fails, just continue without it
            pass
    
    # Data rows
    score_vars = ["Visual Memory", "Verbal Memory", "Visual-Motor Speed", "Reaction Speed"]
    
    # Get both ES and Iverson results regardless of method
    if method == "Individual Outcome Classification (ES Method)":
        # Primary method is ES, but calculate Iverson for comparison
        es_scores = results_data['es_scores']
        adj_scores = results_data['adj_scores']
        
        # Calculate Iverson percentiles for comprehensive report
        try:
            with open('model_parameters.json', 'r') as f:
                params = json.load(f)
            with open('cognitive_models.pkl', 'rb') as f:
                models = pickle.load(f)
            
            percentiles = []
            score_vars_internal = ["visual_memory", "verbal_memory", "visual_motor_speed", "reaction_speed"]
            for i, var in enumerate(score_vars_internal):
                is_reaction_speed = (var == 'reaction_speed')
                percentile = compute_iverson_percentile(
                    patient_data['raw_scores'][i], patient_data['age'], patient_data['sex'], patient_data['education'],
                    models[var], params['sample_stats'][var]['mse_resid'],
                    is_reaction_speed=is_reaction_speed
                )
                percentiles.append(percentile)
            
            iverson_classification = compute_iverson_classification(percentiles)
        except:
            percentiles = [50.0] * 4  # Default if calculation fails
            iverson_classification = "Not Available"
    
    else:
        # Primary method is Iverson, but calculate ES for comparison
        percentiles = results_data['percentiles']
        iverson_classification = results_data['classification']
        
        # Calculate ES scores for comprehensive report
        try:
            with open('model_parameters.json', 'r') as f:
                params = json.load(f)
            with open('cognitive_models.pkl', 'rb') as f:
                models = pickle.load(f)
            
            es_scores = []
            adj_scores = []
            score_vars_internal = ["visual_memory", "verbal_memory", "visual_motor_speed", "reaction_speed"]
            for i, var in enumerate(score_vars_internal):
                is_reaction_speed = (var == 'reaction_speed')
                adj_score = adjust_single_score(
                    patient_data['raw_scores'][i], patient_data['age'], patient_data['sex'], patient_data['education'], 
                    models[var], params['sample_stats'][var]['adj_mean'],
                    is_reaction_speed=is_reaction_speed
                )
                es_score = assign_es_category(adj_score, params['thresholds'][var])
                es_scores.append(es_score)
                adj_scores.append(adj_score)
        except:
            es_scores = [2] * 4  # Default if calculation fails
            adj_scores = [0.0] * 4
    
# Fill data rows
    for i, domain in enumerate(score_vars):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = domain
        
        if domain == "Reaction Speed":
            efficiency = 1/patient_data['raw_scores'][i]
            row_cells[1].text = f'{patient_data["raw_scores"][i]:.3f} (efficiency: {efficiency:.3f})'
        else:
            row_cells[1].text = f'{patient_data["raw_scores"][i]:.0f}'
        
        # Primary metric (ES or Percentile based on method)
        if method == "Individual ES Method":
            row_cells[2].text = f'{adj_scores[i]:.2f}'
        else:
            row_cells[2].text = f'{percentiles[i]:.1f}%'
        
        # ES Classification
        es_score = es_scores[i]
        interpretation = get_es_interpretation(es_score)
        row_cells[3].text = f'ES {es_score} - {interpretation["label"]}'
        
        # Center align all cells
        for cell in row_cells[:4]:  # Only align first 4 cells
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Merge cells in the last column and add Iverson classification
    # Merge cells from row 1 to row 4 (indices 1-4) in column 4 (index 4)
    merged_cell = table.cell(1, 4)
    for row_idx in range(2, 5):  # Merge rows 2, 3, 4 into row 1
        merged_cell = merged_cell.merge(table.cell(row_idx, 4))
    
    # Add the classification text to the merged cell
    merged_cell.text = iverson_classification
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Make the text bold and centered
    for run in merged_cell.paragraphs[0].runs:
        run.font.bold = False
        run.font.size = Pt(12)
    
    # ALWAYS include both ES Table Guide AND Iverson Method Summary
    
    # ES Table Guide
    doc.add_paragraph('ES Table Guide', style='CustomHeader')
    
    try:
        with open('model_parameters.json', 'r') as f:
            params = json.load(f)
        
        es_table = doc.add_table(rows=5, cols=6)
        es_table.style = 'Table Grid'
        
        # Headers with enhanced formatting
        es_headers = es_table.rows[0].cells
        es_headers[0].text = 'Measure'
        es_headers[1].text = 'ES 0'
        es_headers[2].text = 'ES 1'
        es_headers[3].text = 'ES 2'
        es_headers[4].text = 'ES 3'
        es_headers[5].text = 'ES 4'
        
        for cell in es_headers:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(11)
            # Add shading to header using a more compatible method
            try:
                from docx.oxml import parse_xml
                shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
                cell._element.get_or_add_tcPr().append(shading_elm)
            except:
                # If shading fails, just continue without it
                pass
        
        # Data rows with thresholds
        score_vars_internal = ["visual_memory", "verbal_memory", "visual_motor_speed", "reaction_speed"]
        score_labels = ["Visual Memory", "Verbal Memory", "Visual Motor Speed", "Reaction Speed"]
        
        for i, (var, label) in enumerate(zip(score_vars_internal, score_labels)):
            row_cells = es_table.rows[i + 1].cells
            row_cells[0].text = label
            
            if var in params['thresholds']:
                thresholds = params['thresholds'][var]
                row_cells[1].text = f"≤{thresholds['oTL']:.1f}"
                row_cells[2].text = f"{thresholds['oTL']:.1f}-{thresholds['threshold_1']:.1f}"
                row_cells[3].text = f"{thresholds['threshold_1']:.1f}-{thresholds['threshold_2']:.1f}"
                row_cells[4].text = f"{thresholds['threshold_2']:.1f}-{thresholds['median']:.1f}"
                row_cells[5].text = f"≥{thresholds['median']:.1f}"
            
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    except:
        doc.add_paragraph('ES interpretation table unavailable', style='CustomBody')
    
    # Iverson Method Summary
    doc.add_paragraph('Iverson Method Summary', style='CustomHeader')
    
    iverson_table = doc.add_table(rows=6, cols=3)
    iverson_table.style = 'Table Grid'
    
    # Headers with enhanced formatting
    iverson_headers = iverson_table.rows[0].cells
    iverson_headers[0].text = 'Percentile Threshold'
    iverson_headers[1].text = 'Scores below'
    iverson_headers[2].text = 'Measures'
    
    for cell in iverson_headers:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
        # Add shading to header using a more compatible method
        try:
            from docx.oxml import parse_xml
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            cell._element.get_or_add_tcPr().append(shading_elm)
        except:
            # If shading fails, just continue without it
            pass
    
    # Calculate threshold data using available percentiles
    score_names = ["Verbal Memory", "Visual Memory", "Visual-Motor Speed", "Reaction Speed"]
    
    thresholds_list = [2, 5, 10, 16, 25]
    for i, threshold in enumerate(thresholds_list):
        row_cells = iverson_table.rows[i + 1].cells
        row_cells[0].text = f"≤{threshold}{'nd' if threshold == 2 else ('th' if threshold in [5, 25] else 'th')}"
        
        below_count = sum(p <= threshold for p in percentiles)
        row_cells[1].text = str(below_count)
        
        # Name the measures below threshold
        measures_below = [score_names[j] for j, p in enumerate(percentiles) if p <= threshold]
        if measures_below:
            row_cells[2].text = ", ".join(measures_below)
        else:
            row_cells[2].text = "None"
        
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Clinical Summary (based on primary method)
    doc.add_paragraph('Clinical Summary', style='CustomHeader')
    
    if method == "Individual Outcome Classification (ES Method)":
        # Generate ES-based clinical summary
        score_names = ["Verbal Memory", "Visual Memory", "Visual-Motor Speed", "Reaction Speed"]
        
        # Group scores by performance level
        excellent_scores = [score_names[i] for i, es in enumerate(es_scores) if es == 4]
        good_scores = [score_names[i] for i, es in enumerate(es_scores) if es == 3]
        average_scores = [score_names[i] for i, es in enumerate(es_scores) if es == 2]
        below_scores = [score_names[i] for i, es in enumerate(es_scores) if es == 1]
        poor_scores = [score_names[i] for i, es in enumerate(es_scores) if es == 0]
        
        summary_text = "This subject performed "
        
        if excellent_scores:
            summary_text += f"above average (ES 4) on {', '.join(excellent_scores)}. "
        if good_scores:
            summary_text += f"at good average level (ES 3) on {', '.join(good_scores)}. "
        if average_scores:
            summary_text += f"at average level (ES 2) on {', '.join(average_scores)}. "
        if below_scores:
            summary_text += f"below average (ES 1) on {', '.join(below_scores)}. "
        if poor_scores:
            summary_text += f"well below average (ES 0) on {', '.join(poor_scores)}. "
        
        summary_text += f"Overall assessment as per Iverson Schaltz criteria classifies the collective scores as ({iverson_classification})."
        
        doc.add_paragraph(summary_text.strip(), style='CustomBody')
        
    else:
        # Generate Iverson-based clinical summary
        score_names = ["Verbal Memory", "Visual Memory", "Visual-Motor Speed", "Reaction Speed"]
        
        # Group by performance levels
        above_avg = [score_names[i] for i, p in enumerate(percentiles) if p > 50]
        avg_range = [score_names[i] for i, p in enumerate(percentiles) if 25 < p <= 50]
        low_avg = [score_names[i] for i, p in enumerate(percentiles) if 16 < p <= 25]
        below_avg = [score_names[i] for i, p in enumerate(percentiles) if 5 < p <= 16]
        low_range = [score_names[i] for i, p in enumerate(percentiles) if p <= 5]
        
        summary_text = "This subject performed "
        
        if above_avg:
            summary_text += f"above average (>50th percentile) on {', '.join(above_avg)}. "
        if avg_range:
            summary_text += f"in average range (25-50th percentile) on {', '.join(avg_range)}. "
        if low_avg:
            summary_text += f"in low average range (16-25th percentile) on {', '.join(low_avg)}. "
        if below_avg:
            summary_text += f"below average (5-16th percentile) on {', '.join(below_avg)}. "
        if low_range:
            summary_text += f"in low range (≤5th percentile) on {', '.join(low_range)}. "
        
        summary_text += f"Overall assessment as per Iverson Schaltz criteria classifies the collective scores as ({iverson_classification})."
        
        doc.add_paragraph(summary_text.strip(), style='CustomBody')
    
    # Disclaimer
    doc.add_paragraph('Disclaimer', style='CustomHeader')
    disclaimer_text = """This report is generated by an automated normative calculator and should be interpreted by qualified professionals in conjunction with clinical observations and other relevant information. The methodology uses the following references:

1. Aiello EN, Depaoli EG. Norms and standardizations in neuropsychology via equivalent scores: software solutions and practical guides. Neurol Sci. 2022 Feb;43(2):961-966. doi: 10.1007/s10072-021-05374-0. Epub 2021 Jun 17.

2. Iverson, Grant L., and F. M. Webbe. "Evidence-based neuropsychological assessment in sport-related concussion." The handbook of sport neuropsychology (2011): 131-153."""
    
    doc.add_paragraph(disclaimer_text, style='CustomBody')
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer.getvalue() 

def generate_report(method, patient_data, results_data):
    """Generate automated assessment report (legacy markdown version)"""
    # Keep existing markdown function for compatibility
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    report = f"""
# ImPACT Normative Classification Report

**Date:** {timestamp}
**Assessment Method:** {method}

## PATIENT INFORMATION
- **Age:** {patient_data['age']} years
- **Sex:** {'Male' if patient_data['sex'] == 'm' else 'Female'}
- **Education:** {patient_data['education']} years

## RAW TEST SCORES
"""
    
    score_vars = ["Visual Memory", "Verbal Memory", "Visual-Motor Speed", "Reaction Speed"]
    for i, domain in enumerate(score_vars):
        if domain == "Reaction Speed":
            efficiency = 1/patient_data['raw_scores'][i]
            report += f"- **{domain}:** {patient_data['raw_scores'][i]:.3f} seconds (Efficiency: {efficiency:.3f})\n"
        else:
            report += f"- **{domain}:** {patient_data['raw_scores'][i]:.0f}\n"
    
    if method == "Individual Outcome Classification (ES Method)":
        report += f"""
## INDIVIDUAL ES SCORES (0-4 Scale)
"""
        for i, domain in enumerate(score_vars):
            es_score = results_data['es_scores'][i]
            interpretation = get_es_interpretation(es_score)
            report += f"- **{domain}:** ES {es_score} - {interpretation['label']}\n"
        
        report += f"""
## INTERPRETATION
The ES (Estimated Score) method provides individual domain scores on a 0-4 scale after adjusting for demographic factors (age, sex, education). Each domain is evaluated independently based on population percentiles.

**Key Findings:**
- Scores of 0-1 indicate below average performance requiring clinical attention
- Scores of 2-3 indicate average performance
- Scores of 4 indicate above average performance

"""
    
    else:  # Iverson Method
        report += f"""
## INDIVIDUAL DOMAIN PERCENTILES
"""
        for i, domain in enumerate(score_vars):
            percentile = results_data['percentiles'][i]
            if domain == "Reaction Speed":
                efficiency = 1/patient_data['raw_scores'][i]
                report += f"- **{domain}:** {percentile:.1f}th percentile (Efficiency: {efficiency:.3f})\n"
            else:
                report += f"- **{domain}:** {percentile:.1f}th percentile\n"
        
        classification = results_data['classification']
        classification_info = get_iverson_interpretation(classification)
        
        report += f"""
## COLLECTIVE CLASSIFICATION
**Overall Classification:** {classification}

**Clinical Interpretation:** {classification_info['description']}

**Recommended Action:** {classification_info['clinical_note']}

## CLASSIFICATION RATIONALE
The Iverson method considers the pattern of cognitive weaknesses across domains:
"""
        percentiles = results_data['percentiles']
        below_25th = sum(p <= 25 for p in percentiles)
        below_16th = sum(p <= 16 for p in percentiles)
        below_10th = sum(p <= 10 for p in percentiles)
        below_5th = sum(p <= 5 for p in percentiles)
        below_2nd = sum(p <= 2 for p in percentiles)
        
        report += f"- Scores ≤25th percentile: {below_25th}/4 domains\n"
        report += f"- Scores ≤16th percentile: {below_16th}/4 domains\n"
        report += f"- Scores ≤10th percentile: {below_10th}/4 domains\n"
        report += f"- Scores ≤5th percentile: {below_5th}/4 domains\n"
        report += f"- Scores ≤2nd percentile: {below_2nd}/4 domains\n"
    
    report += f"""
## CLINICAL NOTES
- Scores have been adjusted for demographic factors (age, sex, education)
- Consider clinical context, test conditions, and patient history when interpreting results
- Follow-up assessment may be warranted based on clinical judgment

---
*This report was generated by the ImPACT Normative Classification Dashboard*
"""
    
    return report

def main():
    # Load data
    try:
        models, params = load_models_and_parameters()
    except FileNotFoundError:
        st.error("⚠️ Model files not found. Please run fit_models.py first to generate the required model files.")
        st.stop()
    
    # Professional header
    st.markdown("""
    <div class="main-header">
        <h1> ImPACT NORMATIVE ASSESSMENT DASHBOARD</h1>
        
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar for inputs
    with st.sidebar:
        st.header("📋 Assessment Configuration")
        
        # Method selection
        st.subheader("🔬 Classification Method")
        method = st.radio(
            "Choose Assessment Method:",
            ["Individual Outcome Classification (ES Method)", "Collective Classification (Iverson)"],
            help="ES Method: Individual domain scoring (0-4). Iverson Method: Collective classification based on percentile patterns."
        )
        
        st.markdown("---")
        
        # Demographics

        st.subheader("👤 Subject Demographics")
        age = st.number_input(
            "Age (years)", 
            min_value=float(params['demographic_ranges']['age']['min']), 
            max_value=float(params['demographic_ranges']['age']['max']), 
            value=50.0, 
            step=1.0
        )
        
        sex = st.selectbox("Sex", options=['m', 'f'], format_func=lambda x: 'Male' if x == 'm' else 'Female')
        
        # Calculate maximum possible education years based on age (assuming school starts at age 5)
        max_education_for_age = max(0, age - 5)
        
        # Use the lower of the dataset max or age-based max
        effective_max_education = min(
            float(params['demographic_ranges']['education']['max']), 
            max_education_for_age
        )
        
        education = st.number_input(
            "Education (years)", 
            min_value=float(params['demographic_ranges']['education']['min']), 
            max_value=effective_max_education, 
            value=min(12.0, effective_max_education), 
            step=1.0,
            help=f"Maximum education years for age {age}: {max_education_for_age} (assuming school starts at age 5)"
        )
        
        # Additional validation warning
        if education > max_education_for_age:
            st.warning(f"⚠️ Education years ({education}) exceeds maximum possible for age {age} (max: {max_education_for_age} years)")
        elif education == effective_max_education and effective_max_education < float(params['demographic_ranges']['education']['max']):
            st.info(f"Maximum possible years of education for (age {age} are {max_education_for_age} years)")
        
        st.markdown("---")
        
        # Test scores
        st.subheader("🎯 Cognitive Test Scores")
        score_vars = ["visual_memory", "verbal_memory", "visual_motor_speed", "reaction_speed"]
        score_labels = ["Visual Memory", "Verbal Memory", "Visual-Motor Speed", "Reaction Speed"]
        
        raw_scores = {}
        for var, label in zip(score_vars, score_labels):
            if var == "reaction_speed":
                raw_scores[var] = st.number_input(
                    f"{label} (seconds)",
                    min_value=0.3,
                    max_value=2.0,
                    value=1.0,
                    step=0.01,
                    format="%.3f",
                    help="Enter reaction speed in seconds"
                )
            else:
                raw_scores[var] = st.number_input(
                    label,
                    min_value=0.0,
                    max_value=100.0,
                    value=50.0,
                    step=1.0
                )
        
        st.markdown("---")
        
        # Compute button
        compute_pressed = st.button("🔄 COMPUTE ASSESSMENT", type="primary")
    
    # Initialize session state
    if 'computed' not in st.session_state:
        st.session_state.computed = False
        st.session_state.results = None
    
    # Trigger computation
    if compute_pressed or st.session_state.computed:
        st.session_state.computed = True
        
        # Main content area
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.header("📊 Assessment Results")
            
            if method == "Individual Outcome Classification (ES Method)":
                # ES Method computation
                es_scores = []
                adj_scores = []
                
                for var in score_vars:
                    is_reaction_speed = (var == 'reaction_speed')
                    adj_score = adjust_single_score(
                        raw_scores[var], age, sex, education, 
                        models[var], params['sample_stats'][var]['adj_mean'],
                        is_reaction_speed=is_reaction_speed
                    )
                    
                    es_score = assign_es_category(adj_score, params['thresholds'][var])
                    es_scores.append(es_score)
                    adj_scores.append(adj_score)
                
                # Display ES results with gauges
                st.subheader("🎯 Individual ES Scores")
                
                # Create gauge charts in a 2x2 grid
                gauge_cols = st.columns(2)
                for i, (label, es_score, adj_score) in enumerate(zip(score_labels, es_scores, adj_scores)):
                    with gauge_cols[i % 2]:
                        fig = create_es_gauge_chart(es_score, label)
                        st.plotly_chart(fig, use_container_width=True)
                        
                        interpretation = get_es_interpretation(es_score)
                        st.markdown(f"""
                        <div style='background-color: {interpretation['color']}20; padding: 10px; border-radius: 8px; border-left: 4px solid {interpretation['color']}; margin-bottom: 10px; text-align: center;'>
                            <strong>ES {es_score}: {interpretation['label']}</strong><br>
                            <small>{interpretation['description']}</small>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Show adjusted score
                        st.markdown(f"""
                        <div style='text-align: center; margin-bottom: 15px; padding: 8px; background-color: #f0f2f6; border-radius: 5px; border: 1px solid #ddd;'>
                            <strong>Adjusted Score: {adj_score:.3f}</strong>
                        </div>
                        """, unsafe_allow_html=True)
                
                # Store results for report generation
                st.session_state.results = {
                    'method': method,
                    'patient_data': {
                        'age': age, 'sex': sex, 'education': education,
                        'raw_scores': [raw_scores[var] for var in score_vars]
                    },
                    'results_data': {'es_scores': es_scores, 'adj_scores': adj_scores}
                }
                
            else:
                # Iverson Method computation
                percentiles = []
                
                for var in score_vars:
                    is_reaction_speed = (var == 'reaction_speed')
                    percentile = compute_iverson_percentile(
                        raw_scores[var], age, sex, education,
                        models[var], params['sample_stats'][var]['mse_resid'],
                        is_reaction_speed=is_reaction_speed
                    )
                    percentiles.append(percentile)
                
                classification = compute_iverson_classification(percentiles)
                classification_info = get_iverson_interpretation(classification)
                
                # Display percentile results with gauges
                st.subheader("📊 Domain Percentiles")
                
                # Create gauge charts in a 2x2 grid
                gauge_cols = st.columns(2)
                for i, (label, percentile) in enumerate(zip(score_labels, percentiles)):
                    with gauge_cols[i % 2]:
                        fig = create_gauge_chart(percentile, label)
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Show reaction efficiency for reaction speed
                        if label == "Reaction Speed":
                            efficiency = 1/raw_scores['reaction_speed']
                            st.markdown(f"""
                            <div style='text-align: center; margin-bottom: 15px; padding: 5px; background-color: #f0f2f6; border-radius: 5px;'>
                                <small>Efficiency Score: {efficiency:.3f}</small>
                            </div>
                            """, unsafe_allow_html=True)
                
                # Collective Classification
                st.subheader("🎯 Collective Classification")
                st.markdown(f"""
                <div class="classification-card" style='background: linear-gradient(135deg, {classification_info['color']}20, {classification_info['color']}10); border: 2px solid {classification_info['color']};'>
                    <h2 style='color: {classification_info['color']}; margin-bottom: 10px;'>{classification}</h2>

                </div>
                """, unsafe_allow_html=True)
                
                # Store results for report generation
                st.session_state.results = {
                    'method': method,
                    'patient_data': {
                        'age': age, 'sex': sex, 'education': education,
                        'raw_scores': [raw_scores[var] for var in score_vars]
                    },
                    'results_data': {'percentiles': percentiles, 'classification': classification}
                }
        
        with col2:
            st.header("ℹ️ Reference Information")
            
            # Sample information
            with st.expander("📋 Sample Information", expanded=True):
                sample_size = params['sample_stats']['visual_memory']['n_samples']
                age_range = params['demographic_ranges']['age']
                edu_range = params['demographic_ranges']['education']
                
                st.markdown(f"""
                **Normative Sample:**
                - Sample Size: {sample_size:,} participants
                - Age Range: {age_range['min']:.0f} - {age_range['max']:.0f} years
                - Education Range: {edu_range['min']:.0f} - {edu_range['max']:.0f} years
                """)
            
            # Method-specific reference information
            if method == "Individual Outcome Classification (ES Method)":
                with st.expander("📏 Tolerance Limits", expanded=False):
                    for var, label in zip(score_vars, score_labels):
                        thresholds = params['thresholds'][var]
                        
                        # Show statistical tolerance limits information if available
                        if 'oTL_obs' in thresholds:
                            st.markdown(f"""
                            **{label}:**
                            - Outer Tolerance Limit: observation {thresholds['oTL_obs']} (safety level: {thresholds['p_oTL']*100:.1f}%, score ≤{thresholds['oTL']:.2f})
                            - Inner Tolerance Limit: observation {thresholds['iTL_obs']} (safety level: {thresholds['p_iTL']*100:.1f}%, score ≥{thresholds['iTL']:.2f})
                            - ES Thresholds: {thresholds['ES1_obs']}, {thresholds['ES2_obs']}, {thresholds['ES3_obs']} (scores: {thresholds['threshold_1']:.2f}, {thresholds['threshold_2']:.2f}, {thresholds['threshold_3']:.2f})
                            """)
                        else:
                            # Fallback to simple display if statistical limits not available
                            st.markdown(f"""
                            **{label}:**
                            - Lower bound: {thresholds['oTL']:.2f}
                            - Upper bound: {thresholds['iTL']:.2f}
                            - ES cut-offs: {thresholds['threshold_1']:.2f}, {thresholds['threshold_2']:.2f}, {thresholds['threshold_3']:.2f}
                            """)
                
                with st.expander("🎯 ES Score Guide", expanded=False):
                    for i in range(5):
                        interp = get_es_interpretation(i)
                        st.markdown(f"""
                        <div style='background-color: {interp['color']}20; padding: 8px; margin: 4px 0; border-radius: 5px; border-left: 3px solid {interp['color']}; text-align: center;'>
                            <strong>ES {i}: {interp['label']}</strong><br>
                            <small>{interp['description']}</small>
                        </div>
                        """, unsafe_allow_html=True)
            
            else:  # Iverson Method
                # with st.expander("📊 Percentile Thresholds", expanded=True):
                #     st.markdown("""
                #     **Clinical Percentile Thresholds:**
                #     - **≤2nd percentile:** Extremely low performance
                #     - **≤5th percentile:** Very low performance  
                #     - **≤10th percentile:** Low performance
                #     - **≤16th percentile:** Below average performance
                #     - **≤25th percentile:** Low-average performance
                #     - **>25th percentile:** Average or above performance
                #     """)
                
                with st.expander("🏷️ Classification Rules", expanded=False):
                    # Get the colors for each classification
                    broadly_normal_color = get_iverson_interpretation("Broadly Normal")["color"]
                    below_average_color = get_iverson_interpretation("Below Average")["color"]
                    well_below_color = get_iverson_interpretation("Well Below Average")["color"]
                    unusually_low_color = get_iverson_interpretation("Unusually Low")["color"]
                    extremely_low_color = get_iverson_interpretation("Extremely Low")["color"]
                    
                    st.markdown("""**Iverson Classification Criteria:**""")
                    
                    st.markdown(f"""
                    <div style='background-color: {broadly_normal_color}20; padding: 8px; margin: 4px 0; border-radius: 5px; border-left: 3px solid {broadly_normal_color};'>
                        <strong style='color: {broadly_normal_color};'>Broadly Normal:</strong> ≤2 scores ≤25th percentile AND ≤1 score ≤16th percentile AND 0 scores ≤10th percentile
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown(f"""
                    <div style='background-color: {below_average_color}20; padding: 8px; margin: 4px 0; border-radius: 5px; border-left: 3px solid {below_average_color};'>
                        <strong style='color: {below_average_color};'>Below Average:</strong> ≥3 scores ≤25th percentile OR 2 scores ≤16th percentile OR 1 score ≤10th percentile
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown(f"""
                    <div style='background-color: {well_below_color}20; padding: 8px; margin: 4px 0; border-radius: 5px; border-left: 3px solid {well_below_color};'>
                        <strong style='color: {well_below_color};'>Well Below Average:</strong> ≥3 scores ≤16th percentile OR 2 scores ≤10th percentile OR 1 score ≤5th percentile
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown(f"""
                    <div style='background-color: {unusually_low_color}20; padding: 8px; margin: 4px 0; border-radius: 5px; border-left: 3px solid {unusually_low_color};'>
                        <strong style='color: {unusually_low_color};'>Unusually Low:</strong> ≥3 scores ≤10th percentile OR 2 scores ≤5th percentile OR 1 score ≤2nd percentile
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown(f"""
                    <div style='background-color: {extremely_low_color}20; padding: 8px; margin: 4px 0; border-radius: 5px; border-left: 3px solid {extremely_low_color};'>
                        <strong style='color: {extremely_low_color};'>Extremely Low:</strong> ≥3 scores ≤5th percentile OR ≥2 scores ≤2nd percentile
                    </div>
                    """, unsafe_allow_html=True)

            
            # Model statistics
            with st.expander("📊 Model Statistics", expanded=False):
                try:
                    with open('model_summaries.json', 'r') as f:
                        model_summaries = json.load(f)
                    
                    for var, label in zip(score_vars, score_labels):
                        if var in model_summaries:
                            model_info = model_summaries[var]
                            st.markdown(f"""
                            **{label}:**
                            - R²: {model_info['rsquared']:.3f}
                            - Adj. R²: {model_info['rsquared_adj']:.3f}
                            - F-statistic: {model_info['fvalue']:.2f} (p={model_info['f_pvalue']:.3f})
                            """)
                except FileNotFoundError:
                    st.write("Model summary statistics not available.")
            
            # Report generation
            if st.session_state.computed and st.session_state.results:
                st.markdown("---")
                
                if st.button("📄 Generate Word Report", type="secondary"):
                    word_report = generate_word_report(
                        st.session_state.results['method'],
                        st.session_state.results['patient_data'],
                        st.session_state.results['results_data']
                    )
                    
                    st.download_button(
                        label="💾 Download Word Report",
                        data=word_report,
                        file_name=f"cognitive_assessment_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    
    else:
        # Show instructions when not computed
        st.markdown("""
        <div style='background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); padding: 40px; border-radius: 15px; text-align: center; margin: 20px 0;'>
            <h2 style='color: #1e3c72; margin-bottom: 20px;'>🚀 Ready to Begin Assessment</h2>
            <p style='font-size: 18px; color: #2a5298; margin-bottom: 15px;'>Complete the patient information and test scores in the sidebar</p>
            <p style='font-size: 16px; color: #666;'>Then click <strong>"COMPUTE ASSESSMENT"</strong> to generate results</p>
        </div>
        """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()